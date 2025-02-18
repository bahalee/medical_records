from io import BytesIO
from django.db import models
from django.http import HttpResponse
from django.template.loader import render_to_string
from openpyxl import Workbook
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from docx import Document
from docx.shared import Inches
from datetime import datetime
from typing import List
from django.utils.translation import gettext_lazy as _
from django.contrib.auth.models import AbstractBaseUser , PermissionsMixin, BaseUserManager
from django.utils import timezone
class MedecinManager(BaseUserManager):
    def create_user(self, email, password=None, **extra_fields):
        if not email:
            raise ValueError(_('The Email field must be set'))
        email = self.normalize_email(email)
        if 'medico' not in extra_fields:
            from medico.models import Medico
            extra_fields['medico'] = Medico.objects.get(id=1)  
    
        user = self.model(email=email, **extra_fields)
        user.set_password(password)
        user.save(using=self._db)
        return user

    def create_superuser(self, email, password, **extra_fields):
        extra_fields.setdefault("is_staff", True)
        extra_fields.setdefault("is_superuser", True)
    
        if 'medico' not in extra_fields:
            from medico.models import Medico
            extra_fields['medico'] = Medico.objects.get(id=1)  
    
        return self.create_user(email, password, **extra_fields)
class Medico(models.Model):
    id = models.BigAutoField(primary_key=True)
    name = models.CharField(max_length=255)
    specialty = models.CharField(max_length=100)

    def __str__(self):
        return self.name
class Medecin(AbstractBaseUser, PermissionsMixin):
    id = models.BigAutoField(primary_key=True)
    nom_complet = models.CharField(max_length=255)
    email = models.EmailField(_('email address'), unique=True)
    specialite = models.CharField(max_length=100)
    is_active = models.BooleanField(default=True)
    is_staff = models.BooleanField(default=False)
    date_joined = models.DateTimeField(default=timezone.now)
    
    medico = models.ForeignKey(Medico, on_delete=models.CASCADE, null=True, blank=True)

    objects = MedecinManager()
    USERNAME_FIELD = 'email'
    REQUIRED_FIELDS = ['nom_complet', 'specialite']
    groups = models.ManyToManyField(
        'auth.Group',
        related_name='medecin_set',
        blank=True,
        help_text=_('The groups this user belongs to. A user will get all permissions granted to each of their groups.'),
        verbose_name=_('groups'),
    )

    user_permissions = models.ManyToManyField(
        'auth.Permission',
        related_name='medecin_set_permissions',
        blank=True,
        help_text=_('Specific permissions for this user.'),
        verbose_name=_('user permissions'),
    )

    def __str__(self):
        return self.nom_complet

    def gererEnregistrements(self) -> List:
        """Returns the list of enregistrements for the Medecin."""
        return Enregistrement.objects.filter(medecin=self)
class Enregistrement(models.Model):
    idenreg = models.AutoField(primary_key=True)
    nomComplet = models.CharField(max_length=200)
    img1 = models.ImageField(upload_to='uploads/', null=True, blank=True)
    img2 = models.ImageField(upload_to='uploads/', null=True, blank=True)
    img3 = models.ImageField(upload_to='uploads/', null=True, blank=True)
    img4 = models.ImageField(upload_to='uploads/', null=True, blank=True)
    img5 = models.ImageField(upload_to='uploads/', null=True, blank=True)
    naissance = models.DateField()
    admission = models.DateTimeField(default=datetime.now)
    medication = models.TextField()
    notes = models.TextField()
    medecin = models.ForeignKey(Medecin, on_delete=models.CASCADE, null=False)

    def __str__(self):
        return self.nomComplet

    def lister(self) -> List:
        return Enregistrement.objects.filter(medecin=self.medecin)

    def afficher(self, idenreg: int) -> 'Enregistrement':
        return Enregistrement.objects.get(idenreg=idenreg)

    def creer_enreg(self, titre: str, details: str):
        return Enregistrement.objects.create(nomComplet=titre, notes=details, medecin=self)

    def modifier(self, idenreg: int, titre: str, details: str):
        enreg = self.afficher(idenreg)
        enreg.nomComplet = titre
        enreg.notes = details
        enreg.save()

    def supprimer(self, idenreg: int):
        enreg = self.afficher(idenreg)
        enreg.delete()

class Filtrer(models.Model):
    criteres = models.JSONField()

    def appliquerFiltres(self) -> List:
        filters = {key: value for key, value in self.criteres.items()}
        return Enregistrement.objects.filter(**filters)

    def réinitialiserFiltres(self):
        self.criteres = {}
        self.save()
class Exporter(models.Model):
    formatsPrisEnCharge = models.JSONField()
    def exporterEnregistrement(self, enr, format: str) -> HttpResponse:
        if not self.vérifierFormat(format):
            return HttpResponse("Invalid format", status=400)

        export_methods = {
            'XLS': self.export_to_xls,
            'PDF': self.export_to_pdf,
            'DOCX': self.export_to_docx,
            'HTML': self.export_to_html,
        }
        return export_methods[format](enr)

    def vérifierFormat(self, format: str) -> bool:
        return format in self.formatsPrisEnCharge

    def export_to_xls(self, enr) -> HttpResponse:
        wb = Workbook()
        ws = wb.active
        ws.title = "Enregistrement"
        ws.append(["Nom Complet", "Naissance", "Admission", "Medication", "Notes"])
        naissance = enr.naissance
        admission = enr.admission
        if isinstance(naissance, datetime) and naissance.tzinfo:
            naissance = naissance.replace(tzinfo=None) 
        if isinstance(admission, datetime) and admission.tzinfo:
            admission = admission.replace(tzinfo=None)  
        ws.append([enr.nomComplet, naissance, admission, enr.medication, enr.notes])
        response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')
        response['Content-Disposition'] = 'attachment; filename="enregistrement.xlsx"'
        wb.save(response)
        return response
    
    def export_to_pdf(self, enr) -> HttpResponse:
        buffer = BytesIO()
        p = canvas.Canvas(buffer, pagesize=letter)
        p.drawString(100, 750, f"Nom Complet: {enr.nomComplet}")
        p.drawString(100, 730, f"Naissance: {enr.naissance}")
        p.drawString(100, 710, f"Admission: {enr.admission}")
        p.drawString(100, 690, f"Medication: {enr.medication}")
        p.drawString(100, 670, f"Notes: {enr.notes}")
        y_position = 400 
        for i in range(1, 6):
            img_field = getattr(enr, f'img{i}')
            if img_field and img_field.path:
                try:
                    p.drawImage(img_field.path, 100, y_position, width=200, height=200) 
                    y_position -= 210 
                except FileNotFoundError:
                    p.drawString(100, y_position, f"Image {i} not found") 
                    y_position -= 20 
        p.showPage()
        p.save()
        buffer.seek(0)
        response = HttpResponse(buffer, content_type='application/pdf')
        response['Content-Disposition'] = 'attachment; filename="enregistrement.pdf"'
        return response
    def export_to_docx(request, enr):
        """Export the record to a DOCX file."""
        doc = Document() 
        doc.add_heading(f"Enregistrement de {enr.nomComplet}", level=0)
        doc.add_paragraph(f"Nom Complet: {enr.nomComplet}")
        doc.add_paragraph(f"Naissance: {enr.naissance}")
        doc.add_paragraph(f"Admission: {enr.admission}")
        doc.add_paragraph(f"Medication: {enr.medication}")
        doc.add_paragraph(f"Notes: {enr.notes}")
        for i in range(1, 6):
            img_field = getattr(enr, f'img{i}')
            if img_field and img_field.path:
                try:
                    doc.add_picture(img_field.path, width=Inches(2)) 
                    doc.add_paragraph() 
                except FileNotFoundError:
                    doc.add_paragraph(f"Image {i} not found")
        response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        response['Content-Disposition'] = 'attachment; filename="enregistrement.docx"'
        doc.save(response)
        return response
    def export_to_html(self, enr) -> HttpResponse:
        context = {
            'nomComplet': enr.nomComplet,
            'naissance': enr.naissance,
            'admission': enr.admission,
            'medication': enr.medication,
            'notes': enr.notes,
            'enregistrement': enr 
        }
        html_content = render_to_string('medical_records/enregistrement.html', context)
        response = HttpResponse(html_content, content_type='text/html')
        response['Content-Disposition'] = 'attachment; filename="enregistrement.html"'
        return response
from django.db import models
from django.contrib.auth import get_user_model

User = get_user_model()

class Feedback(models.Model):
    user = models.ForeignKey(User, on_delete=models.CASCADE)
    message = models.TextField()
    created_at = models.DateTimeField(auto_now_add=True)

    def __str__(self):
        return f"Feedback from {self.user.email}"